import json 
import openai
from docx import Document
import uuid
import os
from pathlib import Path

# Load environment variables or use a config file instead of hardcoding
# You should set these as environment variables:
# export AZURE_OPENAI_KEY="your_key_here"
# export AZURE_OPENAI_ENDPOINT="your_endpoint_here"
# export AZURE_API_VERSION="2024-05-01-preview"

# Create conversations directory if it doesn't exist
CONVERSATIONS_DIR = Path("conversations")
CONVERSATIONS_DIR.mkdir(exist_ok=True)

try:
    AZURE_CLIENT = openai.AzureOpenAI(
        api_key=AZURE_OPENAI_KEY,
        api_version=AZURE_API_VERSION,
        azure_endpoint=AZURE_OPENAI_ENDPOINT,
    )
except Exception as e:
    print(f"Error initializing Azure OpenAI client: {e}")
    AZURE_CLIENT = None

def read_json_file(file_path="ids.json"):
    """Read and return JSON data from file with error handling."""
    try:
        with open(file_path, "r") as f:
            data = json.load(f)
        return data
    except FileNotFoundError:
        print(f"Error: File '{file_path}' not found.")
        return []
    except json.JSONDecodeError:
        print(f"Error: Invalid JSON in file '{file_path}'.")
        return []
    except Exception as e:
        print(f"Error reading file '{file_path}': {e}")
        return []

def call_openai_history(thread_id):
    """Fetch conversation history and save to Word document."""
    if not AZURE_CLIENT:
        print("Azure OpenAI client not initialized. Skipping...")
        return
    
    try:
        # Fetch messages from the thread
        messages = AZURE_CLIENT.beta.threads.messages.list(thread_id=thread_id)
        messages = messages.data
        
        if not messages:
            print(f"No messages found for thread ID: {thread_id}")
            return
        
        # Reverse messages to show in chronological order
        messages.reverse()
        
        # Create Word document
        doc = Document()
        doc.add_heading(f"Message from conversation Id: {thread_id}", level=0)
        
        for msg in messages:
            if msg.content and len(msg.content) > 0:
                content_text = msg.content[0].text if hasattr(msg.content[0], 'text') else str(msg.content[0])

                if msg.role == "user":
                    p = doc.add_paragraph()
                    run = p.add_run("User: ")
                    run.bold = True
                    p.add_run(content_text.value)
                else:
                    p = doc.add_paragraph()
                    run = p.add_run("Assistant: ")
                    run.bold = True
                    p.add_run(content_text.value)
                
        
        # Generate unique filename and save
        file_name = f"{uuid.uuid4()}.docx"
        file_path = CONVERSATIONS_DIR / file_name
        
        doc.save(str(file_path))
        print(f"Conversation saved to: {file_path}")
        
    except Exception as e:
        print(f"Error processing thread {thread_id}: {e}")

def main():
    """Main function to process all conversation IDs."""
    ids = read_json_file()
    
    if not ids:
        print("No conversation IDs found. Please check your ids.json file.")
        return
    
    print(f"Processing {len(ids)} conversation(s)...")
    
    for id in ids:
        print(f"Processing conversation ID: {id}")
        call_openai_history(id)
    
    print("All conversations processed!")

if __name__ == "__main__":
    main()
    


    





